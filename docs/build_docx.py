#!/usr/bin/env python3
"""
build_docx.py - Sinh file báo cáo Word (.docx) từ docs/BAOCAO.md.

Dùng: python build_docx.py
Yêu cầu: pip install python-docx

Hỗ trợ subset Markdown dùng trong BAOCAO.md: tiêu đề (#..####), đoạn văn,
**đậm** / *nghiêng* / `mã`, khối code ```...```, danh sách -, bảng |...|.
Tự thêm trang bìa và mục lục (placeholder cập nhật trong Word).
"""
import re
import os
import docx
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

HERE = os.path.dirname(os.path.abspath(__file__))
MD = os.path.join(HERE, "BAOCAO.md")
OUT = os.path.join(HERE, "BaoCao_DeTai42_Networking.docx")

# ----- Thông tin trang bìa (chỉnh tại đây) -----
TRUONG = "HỌC VIỆN KỸ THUẬT MẬT MÃ"
KHOA = "KHOA CÔNG NGHỆ THÔNG TIN"
MON = "LẬP TRÌNH NHÂN LINUX"
DE_TAI = ("Đề tài 42: Lập trình shell quản lý hệ thống; quản lý tiến trình, "
          "file, socket & network trong Ubuntu; xây dựng mô-đun nhân tích hợp "
          "vào hệ thống — Networking")
THANH_VIEN = [
    "Ngô Hồng Phong",
    "Ngô Minh Cường",
    "Đỗ Minh Thuần",
]
GVHD = "(Giảng viên hướng dẫn)"
NAM = "Hà Nội, 2026"

MONO = "Consolas"


def set_cell_bg(cell, color):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), color)
    tcPr.append(shd)


def add_inline(paragraph, text):
    """Parse **bold**, *italic*, `code` trong một dòng văn bản."""
    pattern = r"(\*\*.+?\*\*|\*.+?\*|`.+?`)"
    for part in re.split(pattern, text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2]); run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = MONO; run.font.size = Pt(10)
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1]); run.italic = True
        else:
            paragraph.add_run(part)


def add_code_block(doc, lines):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("\n".join(lines))
    run.font.name = MONO
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    # nền xám nhạt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)


def add_table(doc, rows):
    cols = len(rows[0])
    table = doc.add_table(rows=0, cols=cols)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        cells = table.add_row().cells
        for j, val in enumerate(row):
            cells[j].text = ""
            p = cells[j].paragraphs[0]
            add_inline(p, val.strip())
            if i == 0:
                for r in p.runs:
                    r.bold = True
                set_cell_bg(cells[j], "D9E2F3")
    return table


def build_cover(doc):
    def center(text, size, bold=False, color=None, before=0, after=6):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(before)
        p.paragraph_format.space_after = Pt(after)
        r = p.add_run(text); r.font.size = Pt(size); r.bold = bold
        if color:
            r.font.color.rgb = color
        return p

    center(TRUONG, 15, bold=True, before=24)
    center(KHOA, 13, bold=True)
    center("- - - - - o0o - - - - -", 12, after=60)
    center("BÁO CÁO CUỐI KỲ", 20, bold=True, color=RGBColor(0x1F, 0x47, 0x7C))
    center(MON, 16, bold=True, after=40)
    center(DE_TAI, 14, bold=True, after=50)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Sinh viên thực hiện:").bold = True
    for tv in THANH_VIEN:
        center(tv, 13, after=2)
    center("", 8)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Giảng viên hướng dẫn: ").bold = True
    p.add_run(GVHD)
    center(NAM, 13, bold=True, before=50)
    doc.add_page_break()


def build():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(13)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.4

    build_cover(doc)

    # Mục lục (placeholder - cập nhật trong Word: References > Update Table)
    h = doc.add_heading("MỤC LỤC", level=1)
    doc.add_paragraph("(Trong Word: tab References → Table of Contents để chèn "
                      "mục lục tự động từ các tiêu đề.)").italic = True
    doc.add_page_break()

    with open(MD, encoding="utf-8") as f:
        raw = f.read().splitlines()

    # bỏ phần đầu (H1 + đề tài) vì đã có ở trang bìa; bắt đầu từ '## '
    lines, started = [], False
    for ln in raw:
        if not started and ln.startswith("## "):
            started = True
        if started:
            lines.append(ln)

    i = 0
    while i < len(lines):
        ln = lines[i]

        # khối code
        if ln.strip().startswith("```"):
            block = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                block.append(lines[i]); i += 1
            add_code_block(doc, block)
            i += 1
            continue

        # bảng
        if ln.lstrip().startswith("|") and i + 1 < len(lines) \
                and re.match(r"^\s*\|[\s:\-|]+\|\s*$", lines[i + 1]):
            rows = []
            while i < len(lines) and lines[i].lstrip().startswith("|"):
                cells = [c for c in lines[i].strip().strip("|").split("|")]
                if not re.match(r"^[\s:\-|]+$", lines[i].strip().strip("|")):
                    rows.append(cells)
                i += 1
            add_table(doc, rows)
            doc.add_paragraph()
            continue

        # tiêu đề
        m = re.match(r"^(#{2,4})\s+(.*)$", ln)
        if m:
            level = len(m.group(1)) - 1   # ## -> H1
            title = m.group(2).strip().replace("`", "").replace("**", "")
            doc.add_heading(title, level=min(level, 4))
            i += 1
            continue

        # đường kẻ ngang -> bỏ qua
        if ln.strip() == "---":
            i += 1
            continue

        # danh sách
        if re.match(r"^\s*[-*]\s+", ln):
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, re.sub(r"^\s*[-*]\s+", "", ln))
            i += 1
            continue
        if re.match(r"^\s*\d+\.\s+", ln):
            p = doc.add_paragraph(style="List Number")
            add_inline(p, re.sub(r"^\s*\d+\.\s+", "", ln))
            i += 1
            continue

        # đoạn văn thường
        if ln.strip():
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            add_inline(p, ln.strip())
        i += 1

    doc.save(OUT)
    print("Da tao:", OUT)


if __name__ == "__main__":
    build()
